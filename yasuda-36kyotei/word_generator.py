"""
36協定書 Word生成モジュール
python-docxで協定書を動的生成する（5様式パターン対応）
"""
import logging
from typing import Any, Dict, List, Optional, Tuple

from docx import Document
from docx.shared import Pt, Mm, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.table import _Cell
from docx.text.paragraph import Paragraph
from pathlib import Path

# ---------------------------------------------------------------------------
# ロガー設定
# ---------------------------------------------------------------------------
logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# 定数定義
# ---------------------------------------------------------------------------
# フォント
DEFAULT_FONT_NAME: str = "游明朝"

# フォントサイズ (pt)
FONT_SIZE_TITLE: float = 14
FONT_SIZE_SUBTITLE: float = 11
FONT_SIZE_BODY: float = 10.5
FONT_SIZE_ARTICLE: float = 10
FONT_SIZE_TABLE_CELL: float = 9
FONT_SIZE_SEPARATOR: float = 8
FONT_SIZE_NOTE: float = 9
FONT_SIZE_SPECIAL_TITLE: float = 12

# 段落の後方スペース (デフォルト)
DEFAULT_SPACE_AFTER: Pt = Pt(6)
TITLE_SPACE_AFTER: Pt = Pt(12)

# ページ設定 (mm / cm)
PAGE_WIDTH_MM: int = 210
PAGE_HEIGHT_MM: int = 297
MARGIN_CM: float = 2.5

# テーブルスタイル
TABLE_STYLE: str = "Table Grid"

# 区切り線
SEPARATOR_CHAR: str = "─"
SEPARATOR_LENGTH: int = 40
DOUBLE_SEPARATOR_CHAR: str = "="
DOUBLE_SEPARATOR_LENGTH: int = 60


# ---------------------------------------------------------------------------
# ユーティリティ関数
# ---------------------------------------------------------------------------
def set_cell_text(
    cell: _Cell,
    text: str,
    bold: bool = False,
    size: float = FONT_SIZE_ARTICLE,
    align: int = WD_ALIGN_PARAGRAPH.LEFT,
) -> None:
    """セルにテキストを設定"""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(str(text))
    run.font.size = Pt(size)
    run.font.name = DEFAULT_FONT_NAME
    run.bold = bold


def add_paragraph(
    doc: Document,
    text: str,
    bold: bool = False,
    size: float = FONT_SIZE_BODY,
    align: int = WD_ALIGN_PARAGRAPH.LEFT,
    space_after: Pt = DEFAULT_SPACE_AFTER,
) -> Paragraph:
    """段落を追加"""
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = space_after
    run = p.add_run(str(text))
    run.font.size = Pt(size)
    run.font.name = DEFAULT_FONT_NAME
    run.bold = bold
    return p


def _add_table(
    doc: Document,
    rows_data: List[Tuple[str, str]],
) -> None:
    """ラベル-値ペアのリストから2列テーブルを追加する"""
    table = doc.add_table(rows=len(rows_data), cols=2, style=TABLE_STYLE)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value) in enumerate(rows_data):
        set_cell_text(table.cell(i, 0), label, bold=True, size=FONT_SIZE_TABLE_CELL)
        set_cell_text(table.cell(i, 1), value, size=FONT_SIZE_TABLE_CELL)


# ---------------------------------------------------------------------------
# 共通パーツ: ヘッダー（タイトル + 事業所情報 + 前文）
# ---------------------------------------------------------------------------
def _generate_header(
    doc: Document,
    data: Dict[str, Any],
    *,
    subtitle: Optional[str] = None,
    preamble: Optional[str] = None,
) -> None:
    """全様式共通のヘッダー部分を生成する

    Args:
        doc: Wordドキュメント
        data: 入力データ辞書
        subtitle: タイトル直下に表示するサブタイトル（様式9_3〜9_5用）
        preamble: 前文テキスト（Noneの場合は様式9号用のデフォルト前文を使用）
    """
    # タイトル
    add_paragraph(
        doc,
        "時間外労働及び休日労働に関する協定書",
        bold=True,
        size=FONT_SIZE_TITLE,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=TITLE_SPACE_AFTER,
    )

    # サブタイトル（様式9_3〜9_5で使用）
    if subtitle:
        add_paragraph(
            doc, subtitle, size=FONT_SIZE_SUBTITLE, align=WD_ALIGN_PARAGRAPH.CENTER
        )

    # 事業所情報
    add_paragraph(doc, f"事業の名称: {data.get('事業所名', '')}")
    add_paragraph(doc, f"事業の種類: {data.get('事業の種類', '')}")
    # 様式9号のみ電話番号付き、それ以外は事業主名のみ
    if subtitle is None and preamble is None:
        # 様式9号（デフォルト前文）
        add_paragraph(
            doc,
            f"事業主名: {data.get('事業主名', '')}　（電話: {data.get('電話番号', '')}）",
        )
    else:
        add_paragraph(doc, f"事業主名: {data.get('事業主名', '')}")

    # 前文
    add_paragraph(doc, "")
    if preamble is not None:
        add_paragraph(doc, preamble, size=FONT_SIZE_ARTICLE)
    else:
        # 様式9号用デフォルト前文
        add_paragraph(
            doc,
            f"{data.get('事業主名', '')}（以下「甲」という。）と　労働者代表者（以下「乙」という。）は、"
            "労働基準法第３６条第１項の規定に基づき、労働基準法に定める法定労働時間"
            "（１週４０時間、１日８時間）並びに変形労働時間制に定める所定労働時間を超えた労働時間で、"
            "かつ１日８時間、１週４０時間の法定労働時間又は変形期間の法定労働時間の総枠を超える労働"
            "（以下「時間外労働」という。）及び労働基準法に定める休日（毎週１日又は４週４日）における労働"
            "（以下「休日労働」という。）に関し、次の通り協定する。",
            size=FONT_SIZE_ARTICLE,
        )


# ---------------------------------------------------------------------------
# 共通パーツ: 署名欄（フッター）
# ---------------------------------------------------------------------------
def _generate_footer(doc: Document, data: Dict[str, Any]) -> None:
    """全様式共通の署名欄を生成する（旧 _add_signature_section と同一）"""
    add_paragraph(doc, "")
    add_paragraph(doc, SEPARATOR_CHAR * SEPARATOR_LENGTH, size=FONT_SIZE_SEPARATOR)
    add_paragraph(doc, "")

    add_paragraph(doc, f"協定締結日: {data.get('協定締結日', '')}")
    add_paragraph(doc, f"届出作成日: {data.get('届出作成日', '')}")
    add_paragraph(doc, "")
    add_paragraph(
        doc,
        f"（甲）{data.get('事業主職名', '')}　{data.get('届出_事業主名', data.get('事業主名', ''))}",
    )
    add_paragraph(
        doc,
        f"（乙）労働者代表　{data.get('労働者代表_職', '')}　{data.get('労働者代表_氏名', '')}",
    )
    add_paragraph(doc, "")
    add_paragraph(doc, f"所轄労働局: {data.get('所轄労働局', '')}")
    add_paragraph(doc, f"所轄労基署: {data.get('所轄労基署', '')}")


# 後方互換性のためエイリアスを残す
_add_signature_section = _generate_footer


# ---------------------------------------------------------------------------
# 共通パーツ: 簡易時間外労働テーブル（様式9_3/9_4/9_5共通）
# ---------------------------------------------------------------------------
def _generate_simple_overtime_table(
    doc: Document,
    data: Dict[str, Any],
    extra_rows: Optional[List[Tuple[str, str]]] = None,
) -> None:
    """様式9_3/9_4/9_5で共通する時間外労働テーブルを生成する

    Args:
        doc: Wordドキュメント
        data: 入力データ辞書
        extra_rows: テーブル末尾に追加する行（様式固有の項目用）
    """
    add_paragraph(doc, "【時間外労働】", bold=True)
    rows: List[Tuple[str, str]] = [
        ("業務の種類", data.get("時間外_業務の種類", "")),
        ("労働者数", f"{data.get('労働者数', '')}人"),
        ("延長することができる時間（1日）", f"{data.get('延長時間_1日', '')}時間"),
        ("延長することができる時間（1ヶ月）", f"{data.get('延長時間_1ヶ月', '')}時間"),
    ]
    if extra_rows:
        rows.extend(extra_rows)
    _add_table(doc, rows)


# ---------------------------------------------------------------------------
# 各様式の固有部分
# ---------------------------------------------------------------------------
def generate_form_9(doc: Document, data: Dict[str, Any]) -> None:
    """様式第9号: 一般条項（特別条項なし）"""
    # ヘッダー（デフォルト前文を使用）
    _generate_header(doc, data)

    # 第1条〜第3条
    add_paragraph(doc, "")
    add_paragraph(
        doc,
        "第１条　甲は、時間外労働及び休日労働を可能な限り行わせないように努める。",
        bold=True,
        size=FONT_SIZE_ARTICLE,
    )
    add_paragraph(
        doc,
        "第２条　乙は、故意または過失により時間外労働及び休日労働を生じさせない義務を負う。",
        bold=True,
        size=FONT_SIZE_ARTICLE,
    )
    add_paragraph(
        doc,
        "第３条　前2条にも関わらずその必要性を生じた場合、甲は次により時間外労働を行わせることができる。",
        bold=True,
        size=FONT_SIZE_ARTICLE,
    )

    # 時間外労働テーブル
    add_paragraph(doc, "")
    add_paragraph(doc, "【時間外労働】", bold=True)
    _add_table(doc, [
        ("時間外労働をさせる必要のある具体的事由", data.get("時間外_事由", "")),
        ("業務の種類", data.get("時間外_業務の種類", "")),
        ("従事する労働者数（満18歳以上の者）", f"{data.get('労働者数', '')}人"),
        ("所定労働時間", f"{data.get('所定労働時間', '')}時間"),
        ("延長することができる時間（1日）", f"{data.get('延長時間_1日', '')}時間"),
        ("延長することができる時間（1ヶ月）", f"{data.get('延長時間_1ヶ月', '')}時間"),
        ("期間", data.get("時間外_期間", f"令和{data.get('起算日_年', '8')}年{data.get('起算日_月', '4')}月1日から1年間")),
    ])

    # 第4条〜第5条
    add_paragraph(doc, "")
    add_paragraph(
        doc,
        "第４条　甲は、時間外労働を行わせる場合は、原則として、前日の終業時刻までに当該労働者に通知する。",
        size=FONT_SIZE_ARTICLE,
    )
    add_paragraph(
        doc,
        f"第５条　第２条の表における1週、１ヶ月及び１年の起算日はいずれも"
        f"令和{data.get('起算日_年', '8')}年{data.get('起算日_月', '4')}月1日とする。",
        size=FONT_SIZE_ARTICLE,
    )

    # 第6条: 休日労働
    add_paragraph(
        doc,
        "第６条　甲は、必要がある場合には、次により休日労働を行わせることができる。",
        bold=True,
        size=FONT_SIZE_ARTICLE,
    )

    add_paragraph(doc, "")
    add_paragraph(doc, "【休日労働】", bold=True)
    _add_table(doc, [
        ("休日労働をさせる必要のある具体的事由", data.get("休日_事由", "")),
        ("業務の種類", data.get("休日_業務の種類", "")),
        ("所定休日", data.get("所定休日", "")),
        ("労働させることができる休日日数", f"1か月に{data.get('休日労働_日数', '')}日"),
        ("始業及び終業の時刻", data.get("始業終業時刻", "")),
    ])

    # 有効期間
    add_paragraph(doc, "")
    add_paragraph(
        doc,
        f"有効期間: 令和{data.get('起算日_年', '8')}年{data.get('起算日_月', '4')}月1日から1年間",
        bold=True,
    )

    # 署名欄
    _generate_footer(doc, data)


def generate_form_9_2(doc: Document, data: Dict[str, Any]) -> None:
    """様式第9号の2: 特別条項付き"""
    # 一般条項部分を先に生成
    generate_form_9(doc, data)

    # 特別条項セクションを追加
    add_paragraph(doc, "")
    add_paragraph(doc, DOUBLE_SEPARATOR_CHAR * DOUBLE_SEPARATOR_LENGTH, size=FONT_SIZE_SEPARATOR)
    add_paragraph(
        doc,
        "【特別条項】",
        bold=True,
        size=FONT_SIZE_SPECIAL_TITLE,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    add_paragraph(
        doc,
        "一定期間についての延長時間は１か月45時間とする。但し、下記の臨時的事由により"
        "限度時間を超えて労働させる必要がある場合は、労使の協議を経て延長することができる。",
        size=FONT_SIZE_ARTICLE,
    )

    _add_table(doc, [
        ("臨時的に限度時間を超える理由", data.get("特別_理由", "")),
        ("業務の種類", data.get("特別_業務の種類", "")),
        ("労働者数", f"{data.get('特別_労働者数', '')}人"),
        ("限度時間を超えることができる回数", f"{data.get('特別_超過回数', '')}回"),
        ("延長することができる時間数（月）", f"{data.get('特別_延長時間_月', '')}時間"),
        ("延長することができる時間数（年）", f"{data.get('特別_延長時間_年', '')}時間"),
        ("限度時間を超えた場合の割増賃金率", f"{data.get('特別_割増賃金率', '')}％"),
        ("健康及び福祉を確保するための措置", data.get("特別_健康措置_内容", "")),
    ])

    add_paragraph(
        doc,
        f"特別条項の手続き: {data.get('特別_手続き', '労使の協議')}",
        size=FONT_SIZE_ARTICLE,
    )


def generate_form_9_3(doc: Document, data: Dict[str, Any]) -> None:
    """様式第9号の3: 研究開発業務"""
    _generate_header(
        doc,
        data,
        subtitle="（新技術・新商品等の研究開発業務）",
        preamble=(
            "本協定は、労働基準法第36条第11項に基づき、新技術・新商品等の研究開発業務に"
            "従事する労働者に対し、時間外労働及び休日労働を行わせることについて協定する。"
        ),
    )
    add_paragraph(
        doc,
        "※ 研究開発業務は上限規制の適用除外ですが、月100時間超で医師の面接指導が義務付けられます。",
        size=FONT_SIZE_NOTE,
    )

    # 時間外労働テーブル
    _generate_simple_overtime_table(
        doc,
        data,
        extra_rows=[("期間", data.get("時間外_期間", ""))],
    )

    _generate_footer(doc, data)


def generate_form_9_4(doc: Document, data: Dict[str, Any]) -> None:
    """様式第9号の4: 適用猶予事業・業務"""
    _generate_header(
        doc,
        data,
        subtitle="（適用猶予期間中における適用猶予事業・業務）",
        preamble=(
            "本協定は、自動車運転者・建設業・医師等の適用猶予事業・業務に従事する"
            "労働者に対し、時間外労働及び休日労働を行わせることについて協定する。"
        ),
    )

    # 時間外労働テーブル
    _generate_simple_overtime_table(
        doc,
        data,
        extra_rows=[("期間", data.get("時間外_期間", ""))],
    )

    _generate_footer(doc, data)


def generate_form_9_5(doc: Document, data: Dict[str, Any]) -> None:
    """様式第9号の5: 適用猶予＋事業場外みなし労働"""
    _generate_header(
        doc,
        data,
        subtitle="（適用猶予事業・業務 / 事業場外みなし労働時間制）",
        preamble=(
            "本協定は、適用猶予事業・業務において、事業場外労働のみなし労働時間に係る協定の"
            "内容を36協定に付記して届出するものである。"
        ),
    )

    # 時間外労働テーブル（最終行が「事業場外みなし労働時間」になる点が他と異なる）
    _generate_simple_overtime_table(
        doc,
        data,
        extra_rows=[("事業場外みなし労働時間", f"{data.get('所定労働時間', '')}時間")],
    )

    _generate_footer(doc, data)


# ---------------------------------------------------------------------------
# 様式パターン → 生成関数のマッピング
# ---------------------------------------------------------------------------
GENERATORS: Dict[str, Any] = {
    "9":   generate_form_9,
    "9_2": generate_form_9_2,
    "9_3": generate_form_9_3,
    "9_4": generate_form_9_4,
    "9_5": generate_form_9_5,
}

FORM_NAMES: Dict[str, str] = {
    "9":   "様式第9号（一般条項）",
    "9_2": "様式第9号の2（特別条項付き）",
    "9_3": "様式第9号の3（研究開発業務）",
    "9_4": "様式第9号の4（適用猶予事業）",
    "9_5": "様式第9号の5（適用猶予＋事業場外みなし）",
}


def generate_word(data: Dict[str, Any], output_dir: str = "output") -> str:
    """Excelデータ1行分からWord協定書を生成する

    Args:
        data: 入力データ辞書
        output_dir: 出力ディレクトリパス

    Returns:
        生成されたファイルパス
    """
    form_type: str = data.get("様式パターン", "9")
    generator = GENERATORS.get(form_type, generate_form_9)

    doc = Document()

    # ページ設定
    section = doc.sections[0]
    section.page_width = Mm(PAGE_WIDTH_MM)
    section.page_height = Mm(PAGE_HEIGHT_MM)
    section.top_margin = Cm(MARGIN_CM)
    section.bottom_margin = Cm(MARGIN_CM)
    section.left_margin = Cm(MARGIN_CM)
    section.right_margin = Cm(MARGIN_CM)

    # 文書生成
    logger.info("様式 %s の協定書を生成開始: %s", form_type, data.get("事業所名", "不明"))
    generator(doc, data)

    # 保存
    output_path = Path(output_dir)
    output_path.mkdir(parents=True, exist_ok=True)

    事業所名: str = data.get("事業所名", "不明")
    safe_name: str = 事業所名.replace("/", "_").replace("\\", "_").replace(" ", "_")
    filename: str = f"36協定書_{safe_name}_{FORM_NAMES[form_type]}.docx"
    filepath: Path = output_path / filename

    doc.save(str(filepath))
    logger.info("生成完了: %s", filepath)
    return str(filepath)


if __name__ == "__main__":
    # ログ設定（スクリプト直接実行時）
    logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")

    # テスト: サンプルデータで生成
    sample: Dict[str, Any] = {
        "事業所名": "テスト株式会社",
        "事業主名": "山田太郎",
        "電話番号": "03-1234-5678",
        "事業の種類": "情報通信業",
        "時間外_事由": "受注の集中、納期の逼迫",
        "時間外_業務の種類": "システム開発",
        "労働者数": "10",
        "所定労働時間": "8",
        "延長時間_1日": "4",
        "延長時間_1ヶ月": "45",
        "休日_事由": "納期対応",
        "休日_業務の種類": "システム開発",
        "所定休日": "土曜日・日曜日",
        "休日労働_日数": "2",
        "始業終業時刻": "9:00〜18:00",
        "起算日_年": "8",
        "起算日_月": "4",
        "様式パターン": "9",
        "労働者代表_職": "主任",
        "労働者代表_氏名": "鈴木一郎",
        "事業主職名": "代表取締役",
        "所轄労働局": "東京",
        "所轄労基署": "品川",
    }

    path = generate_word(sample, "output")
    logger.info("生成完了: %s", path)
