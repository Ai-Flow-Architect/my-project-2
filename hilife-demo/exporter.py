"""
exporter.py — モニタリング報告書 dict を Word(.docx) または Excel(.xlsx) に書き出す
17項目フォーマット対応・新規作成＆既存ファイル追記に対応
"""
from __future__ import annotations

import io
from datetime import date

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

_HEADER_COLOR = "4A6CF7"
_SECTION_COLOR = "E8ECFF"
_WHITE = "FFFFFF"
_LIGHT_BG = "F4F6FF"


# ══════════════════════════════════════════════════════════
#  ヘルパー
# ══════════════════════════════════════════════════════════

def _fmt_achievement(v) -> str:
    if isinstance(v, dict):
        判定 = v.get("判定", "")
        詳細 = v.get("詳細", "")
        return f"【{判定}】{詳細}" if 判定 else 詳細
    return str(v) if v else ""


def _safe(d: dict, *keys, default="") -> str:
    """ネストしたdictから安全に値を取得"""
    for k in keys:
        if isinstance(d, dict):
            d = d.get(k, default)
        else:
            return default
    return str(d) if d else default


# ══════════════════════════════════════════════════════════
#  Word 出力
# ══════════════════════════════════════════════════════════

def _add_page_break(doc: Document) -> None:
    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    r.append(br)
    p.append(r)
    doc.element.body.append(p)


def _write_monitoring_to_doc(doc: Document, minutes: dict) -> None:
    a = minutes.get("A_基本情報", {})
    b = minutes.get("B_支援計画", {})
    c = minutes.get("C_モニタリング結果", {})

    today = date.today().strftime("%Y年%m月%d日")

    # ── タイトル
    title = doc.add_heading("モニタリング報告書", level=1)
    title.runs[0].font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

    # ── A. 基本情報
    doc.add_heading("A．基本情報", level=2)
    table_a = doc.add_table(rows=5, cols=2)
    table_a.style = "Table Grid"
    rows_a = [
        ("利用者氏名",            a.get("利用者氏名", "")),
        ("モニタリング実施日",    a.get("モニタリング実施日") or today),
        ("作成者氏名（サービス管理責任者）", a.get("作成者氏名", "")),
        ("前回モニタリング実施日", a.get("前回モニタリング実施日", "")),
        ("個別支援計画作成日",    a.get("個別支援計画作成日", "")),
    ]
    for i, (label, val) in enumerate(rows_a):
        table_a.rows[i].cells[0].text = label
        table_a.rows[i].cells[1].text = val
        table_a.rows[i].cells[0].paragraphs[0].runs[0].bold = True
    doc.add_paragraph("")

    # ── B. 支援計画
    doc.add_heading("B．支援計画（個別支援計画からの転記）", level=2)
    for label, val in [
        ("支援の全体方針", b.get("支援の全体方針", "")),
        ("長期目標",       b.get("長期目標", "")),
        ("短期目標",       b.get("短期目標", "")),
    ]:
        p = doc.add_paragraph()
        p.add_run(f"■ {label}　").bold = True
        p.add_run(val or "—")
    doc.add_paragraph("")

    # ── C. モニタリング結果
    doc.add_heading("C．モニタリング結果", level=2)

    achievement = c.get("達成状況の評価", {})
    c_fields = [
        ("全体の状況",            c.get("全体の状況", "")),
        ("本人の感想・満足度",    c.get("本人の感想・満足度", "")),
        ("家族・保護者の意向",    c.get("家族・保護者の意向", "")),
        ("達成状況の評価",        _fmt_achievement(achievement)),
        ("達成されない原因の分析", c.get("達成されない原因の分析", "")),
        ("今後の対応・支援方針",  c.get("今後の対応・支援方針", "")),
        ("計画変更の要否",        c.get("計画変更の要否", "")),
        ("次回モニタリング予定日", c.get("次回モニタリング予定日", "")),
        ("その他留意事項",        c.get("その他留意事項", "特になし")),
    ]
    for label, val in c_fields:
        h = doc.add_paragraph()
        h.add_run(f"■ {label}").bold = True
        body = doc.add_paragraph(val or "—")
        body.runs[0].font.size = Pt(10.5)
    doc.add_paragraph("")

    # ── 同意欄
    doc.add_heading("D．利用者確認・同意", level=2)
    table_d = doc.add_table(rows=2, cols=2)
    table_d.style = "Table Grid"
    table_d.rows[0].cells[0].text = "説明日"
    table_d.rows[0].cells[1].text = ""
    table_d.rows[1].cells[0].text = "利用者署名・押印"
    table_d.rows[1].cells[1].text = ""
    for row in table_d.rows:
        row.cells[0].paragraphs[0].runs[0].bold = True


def to_word(minutes: dict, existing_bytes: bytes | None = None) -> bytes:
    if existing_bytes:
        doc = Document(io.BytesIO(existing_bytes))
        _add_page_break(doc)
    else:
        doc = Document()
    _write_monitoring_to_doc(doc, minutes)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ══════════════════════════════════════════════════════════
#  Excel 出力
# ══════════════════════════════════════════════════════════

def _thin_border() -> Border:
    s = Side(style="thin", color="CCCCCC")
    return Border(left=s, right=s, top=s, bottom=s)


def _section_header(ws, row: int, text: str) -> None:
    c = ws.cell(row=row, column=1, value=text)
    ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=2)
    c.font = Font(bold=True, color="FFFFFF", size=11)
    c.fill = PatternFill("solid", fgColor=_HEADER_COLOR)
    c.alignment = Alignment(horizontal="left", vertical="center", indent=1)
    ws.row_dimensions[row].height = 22


def _data_row(ws, row: int, label: str, value: str, bg: bool = False, height: int = 15) -> None:
    c1 = ws.cell(row=row, column=1, value=label)
    c2 = ws.cell(row=row, column=2, value=value)
    c1.font = Font(bold=True, size=10)
    for c in (c1, c2):
        c.border = _thin_border()
        c.alignment = Alignment(vertical="top", wrap_text=True)
        c.fill = PatternFill("solid", fgColor=_LIGHT_BG if bg else _WHITE)
    c2.font = Font(size=10)
    ws.row_dimensions[row].height = height


def _append_monitoring_to_sheet(ws, minutes: dict, start_row: int) -> int:
    ws.column_dimensions["A"].width = 26
    ws.column_dimensions["B"].width = 62

    a = minutes.get("A_基本情報", {})
    b = minutes.get("B_支援計画", {})
    c = minutes.get("C_モニタリング結果", {})
    today = date.today().strftime("%Y年%m月%d日")

    row = start_row

    # ── A. 基本情報
    _section_header(ws, row, "A．基本情報"); row += 1
    _data_row(ws, row, "利用者氏名",   a.get("利用者氏名", ""));  row += 1
    _data_row(ws, row, "モニタリング実施日", a.get("モニタリング実施日") or today, bg=True); row += 1
    _data_row(ws, row, "作成者氏名（サービス管理責任者）", a.get("作成者氏名", "")); row += 1
    _data_row(ws, row, "前回モニタリング実施日", a.get("前回モニタリング実施日", ""), bg=True); row += 1
    _data_row(ws, row, "個別支援計画作成日", a.get("個別支援計画作成日", "")); row += 2

    # ── B. 支援計画
    _section_header(ws, row, "B．支援計画（個別支援計画からの転記）"); row += 1
    _data_row(ws, row, "支援の全体方針", b.get("支援の全体方針", ""), height=40); row += 1
    _data_row(ws, row, "長期目標",       b.get("長期目標", ""),       bg=True, height=30); row += 1
    _data_row(ws, row, "短期目標",       b.get("短期目標", ""),       height=30); row += 2

    # ── C. モニタリング結果
    _section_header(ws, row, "C．モニタリング結果"); row += 1
    achievement = c.get("達成状況の評価", {})
    c_fields = [
        ("全体の状況",            c.get("全体の状況", ""),             False, 60),
        ("本人の感想・満足度",    c.get("本人の感想・満足度", ""),     True,  50),
        ("家族・保護者の意向",    c.get("家族・保護者の意向", ""),     False, 30),
        ("達成状況の評価",        _fmt_achievement(achievement),       True,  40),
        ("達成されない原因の分析", c.get("達成されない原因の分析", ""), False, 40),
        ("今後の対応・支援方針",  c.get("今後の対応・支援方針", ""),   True,  50),
        ("計画変更の要否",        c.get("計画変更の要否", ""),         False, 20),
        ("次回モニタリング予定日", c.get("次回モニタリング予定日", ""), True,  20),
        ("その他留意事項",        c.get("その他留意事項", "特になし"), False, 30),
    ]
    for label, val, bg, h in c_fields:
        _data_row(ws, row, label, val, bg=bg, height=h); row += 1
    row += 1

    # ── D. 同意欄
    _section_header(ws, row, "D．利用者確認・同意"); row += 1
    _data_row(ws, row, "説明日",         ""); row += 1
    _data_row(ws, row, "利用者署名・押印", "", height=30); row += 2

    return row


def to_excel(minutes: dict, existing_bytes: bytes | None = None) -> bytes:
    if existing_bytes:
        wb = openpyxl.load_workbook(io.BytesIO(existing_bytes))
        ws = wb.active
        start_row = ws.max_row + 3
    else:
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "モニタリング報告書"
        start_row = 1

    _append_monitoring_to_sheet(ws, minutes, start_row)
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()
